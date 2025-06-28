import flet as ft
import os
import sys
import time
import threading
from pathlib import Path
from docx import Document
from sentence_transformers import SentenceTransformer

# --- Configuração de Paths ---
def get_base_path():
    """Retorna o caminho base correto em dev e em modo compilado."""
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent
    else:
        return Path(__file__).resolve().parent

BASE_DIR = get_base_path()
ASSETS_DIR = BASE_DIR / "assets"
TEMP_EXPORTS_DIR = ASSETS_DIR / "temp_exports"
TEMP_EXPORTS_URL_PATH = "temp_exports"
MODEL_PATH = ASSETS_DIR / "models" / "all-MiniLM-L6-v2"

# Garante que os diretórios existam
TEMP_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

# --- Variáveis Globais para o Modelo ---
st_model = None
model_ready_event = threading.Event()

def load_model_in_background(model_path: Path):
    """
    Carrega o modelo SentenceTransformer em uma thread de background
    para não bloquear a interface principal do Flet.
    """
    global st_model
    print(f"Iniciando carregamento do modelo de '{model_path}' em background...")
    try:
        if not model_path.exists():
            raise FileNotFoundError(f"Diretório do modelo não encontrado: {model_path}")
        st_model = SentenceTransformer(str(model_path))
        print("Modelo carregado com sucesso!")
    except Exception as e:
        print(f"ERRO CRÍTICO ao carregar o modelo: {e}")
        st_model = None # Garante que o modelo seja None em caso de erro
    finally:
        # Sinaliza que o processo de carregamento (com sucesso ou falha) terminou.
        model_ready_event.set()

class SimpleDocxExporter:
    """Versão simplificada do DocxExporter para este teste."""
    def export_simple_docx(self, data: dict, output_path: str) -> bool:
        try:
            document = Document()
            document.add_heading(data.get("title", "Relatório de Teste"), level=1)
            document.add_paragraph(f"Conteúdo: {data.get('content', '')}")
            if "embedding_shape" in data:
                document.add_paragraph(f"Shape do Embedding: {data.get('embedding_shape')}")
            document.save(output_path)
            return True
        except Exception as e:
            print(f"Erro ao gerar DOCX: {e}")
            return False

def main(page: ft.Page):
    page.title = "Teste de Exportação e SentenceTransformer"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    exporter = SimpleDocxExporter()

    # --- Controles da UI ---
    model_status_text = ft.Text("Carregando modelo de IA em segundo plano...", italic=True)
    text_to_embed = ft.TextField(label="Texto para gerar embedding", value="Este é um teste da biblioteca Sentence Transformers.")
    embedding_result_text = ft.Text(visible=False)
    docx_loading_status = ft.Text(visible=False, text_align=ft.TextAlign.CENTER)

    def update_model_status_on_load():
        """Atualiza a UI quando o modelo estiver pronto."""
        model_ready_event.wait() # Espera o evento
        if st_model:
            model_status_text.value = "Modelo de IA carregado e pronto!"
            model_status_text.color = ft.Colors.GREEN
            embed_button.disabled = False
        else:
            model_status_text.value = "Falha ao carregar o modelo de IA."
            model_status_text.color = ft.Colors.RED
            embed_button.disabled = True
        
        # Como esta função roda em uma thread, precisamos usar page.update() para refletir a mudança.
        page.update()

    def generate_embedding(e):
        """Gera o embedding para o texto de entrada."""
        # A verificação na UI (botão desabilitado) já previne isso, mas é uma boa prática.
        if not st_model:
            embedding_result_text.value = "Modelo não está disponível."
            embedding_result_text.visible = True
            page.update()
            return

        input_text = text_to_embed.value
        if not input_text:
            embedding_result_text.value = "Por favor, insira um texto."
            embedding_result_text.visible = True
            page.update()
            return
            
        embedding_result_text.value = "Gerando embedding..."
        embedding_result_text.visible = True
        page.update()

        try:
            embedding = st_model.encode([input_text])
            embedding_result_text.value = f"Embedding gerado! Shape: {embedding.shape}"
            embedding_result_text.color = ft.Colors.BLUE
        except Exception as ex:
            embedding_result_text.value = f"Erro ao gerar embedding: {ex}"
            embedding_result_text.color = ft.Colors.RED

        embedding_result_text.visible = True
        page.update()

    def generate_and_export_docx(e):
        """Gera o .docx, incluindo o embedding se disponível."""
        docx_loading_status.value = "Gerando arquivo .docx..."
        docx_loading_status.visible = True
        page.update()

        report_data = {
            "title": "Relatório de Teste de Exportação com Embedding",
            "content": text_to_embed.value
        }
        
        # Adiciona o shape do embedding ao relatório se ele foi gerado
        if "Shape" in embedding_result_text.value:
            report_data["embedding_shape"] = embedding_result_text.value.split("Shape: ")[1]

        timestamp = int(time.time())
        filename = f"relatorio_completo_{timestamp}.docx"
        server_file_path = TEMP_EXPORTS_DIR / filename
        url_file_path = f"{TEMP_EXPORTS_URL_PATH}/{filename}"

        success = exporter.export_simple_docx(report_data, str(server_file_path))

        if success:
            docx_loading_status.value = f"Arquivo '{filename}' gerado. Iniciando download..."
            page.update()
            page.launch_url(url_file_path, web_window_name="_blank")
            time.sleep(3)
            docx_loading_status.visible = False
        else:
            docx_loading_status.value = "Erro ao gerar o arquivo .docx."
            docx_loading_status.color = ft.Colors.RED
        page.update()

    embed_button = ft.ElevatedButton("Gerar Embedding", icon=ft.Icons.MEMORY, on_click=generate_embedding, disabled=True)
    export_button = ft.ElevatedButton("Gerar e Exportar DOCX", icon=ft.Icons.FILE_DOWNLOAD, on_click=generate_and_export_docx)

    page.add(
        ft.Column(
            [
                ft.Text("Teste com SentenceTransformer", style=ft.TextThemeStyle.HEADLINE_MEDIUM),
                model_status_text,
                ft.Divider(),
                text_to_embed,
                embed_button,
                embedding_result_text,
                ft.Divider(),
                export_button,
                docx_loading_status,
            ],
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
            spacing=15,
            width=600
        )
    )
    
    # Inicia o carregamento do modelo em background assim que a UI é montada
    threading.Thread(target=update_model_status_on_load, daemon=True).start()

# Inicia a thread que carrega o modelo em background.
# É importante que isso aconteça apenas uma vez.
threading.Thread(target=load_model_in_background, args=(MODEL_PATH,), daemon=True).start()

if __name__ == "__main__":
    ft.app(
        target=main,
        view=ft.AppView.WEB_BROWSER,
        assets_dir=str(ASSETS_DIR)
    )
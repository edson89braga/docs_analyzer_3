# src/core/doc_generator.py
import os
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil # Para copiar templates

from src.core.prompts import formatted_initial_analysis # Para type hinting e acesso aos campos
from src.settings import ASSETS_DIR_ABS # Para acessar a pasta de assets

from src.logger.logger import LoggerSetup
_logger = LoggerSetup.get_logger(__name__)

TEMPLATES_SUBDIR = "templates_docx"

class DocxExporter:
    def __init__(self):
        self.templates_dir = os.path.join(ASSETS_DIR_ABS, TEMPLATES_SUBDIR)
        os.makedirs(self.templates_dir, exist_ok=True) # Garante que o diretório exista

    def _get_field_display_name(self, field_name: str) -> str:
        """Retorna um nome amigável para o campo (para cabeçalhos de tabela, etc.)."""
        # Este mapeamento pode ser expandido ou melhorado
        name_map = {
            "descricao_geral": "Descrição Geral",
            "tipo_documento_origem": "Tipo do Documento de Origem",
            "orgao_origem": "Órgão de Origem",
            "uf_origem": "UF de Origem",
            "municipio_origem": "Município de Origem",
            "resumo_fato": "Resumo do Fato",
            "tipo_local": "Tipo do Local do Fato",
            "uf_fato": "UF do Fato",
            "municipio_fato": "Município do Fato",
            "valor_apuracao": "Valor da Apuração (R$)",
            "area_atribuicao": "Área de Atribuição",
            "tipificacao_penal": "Tipificação Penal",
            "tipo_a_autuar": "Tipo a Autuar",
            "assunto_re": "Assunto (RE)",
            "destinacao": "Destinação",
            "pessoas_envolvidas": "Pessoas Envolvidas",
            "linha_do_tempo": "Linha do Tempo",
            "observacoes": "Observações"
        }
        return name_map.get(field_name, field_name.replace("_", " ").title())

    def export_simple_docx(self, data: formatted_initial_analysis, output_path: str) -> bool:
        """
        Exporta os dados da análise (objeto FormatAnaliseInicial) para uma tabela em um arquivo DOCX.

        Args:
            data (FormatAnaliseInicial): O objeto contendo os dados da análise.
            output_path (str): O caminho completo onde o arquivo .docx será salvo.

        Returns:
            bool: True se a exportação for bem-sucedida, False caso contrário.
        """
        _logger.info(f"Iniciando exportação simples para DOCX em: {output_path}")
        try:
            document = Document()
            document.add_heading('Relatório de Análise do Documento', level=1)

            # Campos a serem incluídos na tabela (excluindo justificativas)
            # E também excluindo os campos multiline que serão tratados separadamente.
            fields_for_table = [
                "tipo_documento_origem", "orgao_origem", "uf_origem", "municipio_origem",
                "tipo_local", "uf_fato", "municipio_fato", "valor_apuracao",
                "area_atribuicao", "tipificacao_penal", "tipo_a_autuar", "assunto_re", "destinacao"
            ]
            
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid' # Estilo de tabela básico
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Campo'
            hdr_cells[1].text = 'Valor'
            # Deixar as células do cabeçalho em negrito
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            for field_name in fields_for_table:
                if hasattr(data, field_name):
                    value = getattr(data, field_name)
                    display_name = self._get_field_display_name(field_name)
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = display_name
                    if isinstance(value, list): # Para campos que podem ser listas (embora aqui sejam strings)
                        row_cells[1].text = ", ".join(map(str, value)) if value else "N/A"
                    elif isinstance(value, float) and field_name == "valor_apuracao":
                        row_cells[1].text = f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    else:
                        row_cells[1].text = str(value) if value is not None else "N/A"
            
            # Adicionando campos multiline separadamente, abaixo da tabela
            multiline_fields = {
                "descricao_geral": "Descrição Geral do Documento",
                "resumo_fato": "Resumo do Fato",
                "pessoas_envolvidas": "Pessoas Envolvidas", # Será uma lista de strings
                "linha_do_tempo": "Linha do Tempo",       # Será uma lista de strings
                "observacoes": "Observações Adicionais"
            }

            for field_name, heading_text in multiline_fields.items():
                if hasattr(data, field_name):
                    value = getattr(data, field_name)
                    if value: # Só adiciona se tiver conteúdo
                        document.add_heading(heading_text, level=2)
                        if isinstance(value, list):
                            for item in value:
                                document.add_paragraph(str(item), style='ListBullet') # Ou 'ListNumber'
                        else:
                            document.add_paragraph(str(value))
                        document.add_paragraph() # Espaço extra

            document.save(output_path)
            _logger.info(f"DOCX simples salvo com sucesso em: {output_path}")
            return True
        except Exception as e:
            _logger.error(f"Erro ao exportar DOCX simples para '{output_path}': {e}", exc_info=True)
            return False

    def get_available_templates(self) -> List[Tuple[str, str]]:
        """
        Lista os arquivos de template .docx disponíveis na pasta de templates.

        Returns:
            List[Tuple[str, str]]: Lista de tuplas, onde cada tupla contém
                                   (nome_amigavel_do_template, caminho_completo_do_template).
        """
        templates = []
        if not os.path.isdir(self.templates_dir):
            _logger.warning(f"Diretório de templates '{self.templates_dir}' não encontrado.")
            return templates

        for filename in os.listdir(self.templates_dir):
            if filename.lower().endswith(".docx"):
                full_path = os.path.join(self.templates_dir, filename)
                # Cria um nome amigável removendo o prefixo e a extensão
                friendly_name = filename[:-len(".docx")].replace("_", " ").title()
                templates.append((friendly_name, full_path))
        
        _logger.info(f"Encontrados {len(templates)} templates em '{self.templates_dir}'.")
        return templates

    def export_from_template_docx(self, data: formatted_initial_analysis, template_path: str, output_path: str) -> Tuple[bool, List[str]]:
        """
        Cria um novo DOCX a partir de um template, substituindo placeholders pelos dados da análise.

        Args:
            data (FormatAnaliseInicial): O objeto contendo os dados da análise.
            template_path (str): Caminho para o arquivo .docx do template.
            output_path (str): Caminho completo onde o arquivo .docx final será salvo.

        Returns:
            Tuple[bool, List[str]]: (True se sucesso False caso contrário, Lista de chaves não encontradas no template mas com valor nos dados)
        """
        _logger.info(f"Iniciando exportação baseada no template '{template_path}' para '{output_path}'")
        missing_keys_with_values: List[str] = []
        try:
            if not os.path.exists(template_path):
                _logger.error(f"Template '{template_path}' não encontrado.")
                return False, ["Template não encontrado"]

            # Copia o template para o local de saída para não modificar o original
            # shutil.copyfile(template_path, output_path) # Cuidado, Docx não gosta de abrir um arquivo que ele está editando
            document = Document(template_path) # Abre o template diretamente

            # Placeholders a serem substituídos (devem corresponder aos campos de FormatAnaliseInicial)
            # Exclui as justificativas e campos que podem não fazer sentido em todos os templates
            fields_to_replace = [
                "tipo_documento_origem", "orgao_origem", "uf_origem", "municipio_origem",
                "tipo_local", "uf_fato", "municipio_fato", "valor_apuracao",
                "area_atribuicao", "tipificacao_penal", "tipo_a_autuar", "assunto_re", "destinacao",
                "descricao_geral", "resumo_fato", "pessoas_envolvidas", "linha_do_tempo", "observacoes"
            ]

            placeholders_found_in_template = set()

            # Função auxiliar para substituir texto em parágrafos e tabelas
            # def replace_text_in_element(element, placeholder, replacement_text):
            #    if hasattr(element, 'text'):
            #        if placeholder in element.text:
            #            placeholders_found_in_template.add(placeholder.strip("<>"))
            #            # A substituição em 'runs' é mais robusta para manter a formatação
            #            for run in element.runs:
            #                if placeholder in run.text:
            #                    run.text = run.text.replace(placeholder, replacement_text)
            #    # Se for um parágrafo, verifica os 'runs'
            #    if hasattr(element, 'paragraphs'):
            #        for p in element.paragraphs:
            #            replace_text_in_element(p, placeholder, replacement_text)

            def replace_text_in_element(paragraph, field_name, value):
                placeholder = f"<{field_name}>" # Formato do placeholder: <nome_do_campo>
                placeholder_found = False
                inline = paragraph.runs
                
                # Primeira tentativa: substituição direta em runs individuais
                for line in range(len(inline)):
                    if inline[line].text != '' and placeholder in inline[line].text:
                        placeholders_found_in_template.add(placeholder[1:-1])
                        inline[line].text = inline[line].text.replace(placeholder, str(value))
                        placeholder_found = True
                
                # Se não encontrou o placeholder em runs individuais
                if not placeholder_found:
                    # Tenta encontrar placeholders que atravessam múltiplos runs
                    # Encontra os runs que contêm partes do placeholder
                    current_text = ''
                    start_run_index = -1
                    last_run_index = -1
                    
                    # Percorre todos os runs para encontrar o placeholder completo
                    for run_index, run in enumerate(inline):
                        current_text += run.text
                        
                        # Verifica se o placeholder foi completamente encontrado
                        if placeholder in current_text:
                            placeholders_found_in_template.add(placeholder[1:-1])
                            last_run_index = run_index
                            
                            # Encontra o primeiro run do placeholder
                            temp_text = ''
                            for i in range(run_index, -1, -1):
                                temp_text = inline[i].text + temp_text
                                if placeholder in temp_text:
                                    start_run_index = i
                                    break
                            
                            # Preserva formatação do primeiro run
                            first_run = inline[start_run_index]
                            
                            # Calcula o texto total dos runs do placeholder
                            total_text = ''.join(r.text for r in inline[start_run_index:last_run_index+1])
                            total_text = total_text.replace(placeholder, str(value))
                            
                            # Remove runs extras
                            for extra_run in inline[start_run_index+1:last_run_index+1]:
                                extra_run.text = ''
                            
                            # Substitui no primeiro run
                            first_run.text = total_text
                            
                            placeholder_found = True
                            break
                        
                        if placeholder_found:
                            break

            for field_name in fields_to_replace:
                if hasattr(data, field_name):
                    value = getattr(data, field_name)
                    
                    # Formata o valor para string
                    if isinstance(value, list):
                        replacement = "\n".join(map(str, value)) if value else "" # Lista de itens, um por linha
                    elif isinstance(value, float) and field_name == "valor_apuracao":
                        replacement = f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if value is not None else ""
                    else:
                        replacement = str(value) if value is not None else ""

                    # Substitui em todos os parágrafos
                    for paragraph in document.paragraphs:
                        replace_text_in_element(paragraph, field_name, replacement)

                    # Substitui em todas as tabelas
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_text_in_element(paragraph, field_name, replacement)
            
            # Verifica por placeholders não encontrados no template mas com valor nos dados
            for field_name in fields_to_replace:
                value = getattr(data, field_name, None)
                if value and field_name not in placeholders_found_in_template:
                    # Verifica se o valor não é apenas uma lista vazia ou string vazia
                    if isinstance(value, list) and not any(value): continue
                    if isinstance(value, str) and not value.strip(): continue
                    
                    missing_keys_with_values.append(self._get_field_display_name(field_name))


            document.save(output_path)
            _logger.info(f"DOCX a partir de template salvo com sucesso em: {output_path}")
            if missing_keys_with_values:
                 _logger.warning(f"Algumas chaves com valor nos dados não foram encontradas como placeholders no template: {missing_keys_with_values}")
            return True, missing_keys_with_values

        except Exception as e:
            _logger.error(f"Erro ao exportar DOCX de template '{template_path}' para '{output_path}': {e}", exc_info=True)
            return False, ["Erro interno durante a exportação."]


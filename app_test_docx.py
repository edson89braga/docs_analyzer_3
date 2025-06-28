import flet as ft
import os, sys
import time
from docx import Document

# --- Configuração de Paths para o Teste (Versão Robusta) ---
SRC_DIR = os.path.abspath(__file__)

def get_resource_path(relative_path=''):
    """
    Obtém o caminho correto para recursos/assets tanto em desenvolvimento
    quanto em ambiente frozen
    """
    if getattr(sys, 'frozen', False):
        # PyInstaller cria uma pasta temporária
        base_path = os.path.dirname(sys.executable)
    else:
        # Desenvolvimento
        base_path = os.path.dirname(SRC_DIR)
    
    return os.path.join(base_path, relative_path)

BASE_DIR = get_resource_path()
ASSETS_DEMO_DIR = os.path.join(BASE_DIR, "assets_demo")
TEMP_EXPORTS_DIR = os.path.join(ASSETS_DEMO_DIR, "temp_exports")
TEMP_EXPORTS_URL_PATH = "temp_exports" # Path relativo para a URL

# Garante que os diretórios existam
from pathlib import Path

TEMP_EXPORTS_DIR = Path(TEMP_EXPORTS_DIR)
TEMP_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

class SimpleDocxExporter:
    """Versão simplificada do DocxExporter para este teste."""
    def export_simple_docx(self, data: dict, output_path: str) -> bool:
        try:
            document = Document()
            document.add_heading(data.get("title", "Relatório de Teste"), level=1)
            document.add_paragraph(data.get("content", "Nenhum conteúdo fornecido."))
            p = document.add_paragraph('Um parágrafo com um pouco de ')
            p.add_run('texto em negrito').bold = True
            p.add_run(' e um pouco de ')
            p.add_run('texto em itálico.').italic = True
            document.save(output_path)
            return True
        except Exception as e:
            print(f"Erro ao gerar DOCX: {e}")
            return False

def main(page: ft.Page):
    page.title = "Teste de Exportação DOCX"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    exporter = SimpleDocxExporter()

    def generate_and_export_docx(e):
        """Gera o arquivo .docx no servidor e tenta iniciar o download."""
        loading_status.value = "Gerando arquivo .docx..."
        loading_status.visible = True
        page.update()

        # Dados de exemplo para o documento
        report_data = {
            "title": "Relatório de Teste de Exportação",
            "content": f"Este documento foi gerado em {time.strftime('%d/%m/%Y %H:%M:%S')}."
        }

        # Gera um nome de arquivo único para evitar conflitos
        timestamp = int(time.time())
        filename = f"relatorio_teste_{timestamp}.docx"
        
        # Caminho completo no sistema de arquivos do servidor
        server_file_path = TEMP_EXPORTS_DIR / filename
        
        # Caminho relativo para ser usado na URL
        # Ex: "temp_exports/relatorio_teste_12345.docx"
        url_file_path = f"{TEMP_EXPORTS_URL_PATH}/{filename}"

        success = exporter.export_simple_docx(report_data, str(server_file_path))

        if success:
            loading_status.value = f"Arquivo '{filename}' gerado. Tentando iniciar download..."
            page.update()
            
            # Tenta iniciar o download usando a URL relativa.
            page.launch_url(url_file_path, web_window_name="_blank")
            
            # Delay para o usuário ver a mensagem antes de escondê-la
            time.sleep(3)
            loading_status.visible = False
            page.update()
        else:
            loading_status.value = "Erro ao gerar o arquivo .docx."
            loading_status.color = ft.Colors.RED
            page.update()

    loading_status = ft.Text(visible=False, text_align=ft.TextAlign.CENTER)
    export_button = ft.ElevatedButton(
        "Gerar e Exportar DOCX",
        icon=ft.Icons.FILE_DOWNLOAD,
        on_click=generate_and_export_docx
    )

    page.add(
        ft.Column(
            [
                ft.Text("Teste de Exportação DOCX em Flet Web", style=ft.TextThemeStyle.HEADLINE_MEDIUM),
                ft.Text("Clique no botão para gerar um arquivo .docx e iniciar o download."),
                ft.Container(height=20),
                export_button,
                ft.Container(height=20),
                loading_status
            ],
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
            spacing=10
        )
    )

if __name__ == "__main__":
    ft.app(
        target=main,
        view=ft.AppView.WEB_BROWSER, 
        port=8550,                   
        assets_dir=ASSETS_DEMO_DIR, 
    )